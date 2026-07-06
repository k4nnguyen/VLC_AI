# Loader cho cac file docx
from pathlib import Path
from docx import Document
from src.loaders.base_loader import BaseLoader
from src.core.models.raw_document import RawDocument

class DocxLoader(BaseLoader):
    def load(self, source: Path) -> RawDocument:
        document = Document(source)
        paragraphs = []
        
        for p in document.paragraphs:
            text = p.text.strip()
            if text:
                paragraphs.append(text)
        
        return RawDocument(
            source= source,
            filename= source.name,
            file_type= "docx",
            raw_text= "\n".join(paragraphs),
        )